"""
Программный генератор Word-шаблона договора.
Создает template_prepared.docx без необходимости исходного template.docx
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from .constants import (
    TEMPLATE_PATH,
    EXECUTOR_DATA,
    EXECUTOR_PREAMBLE,
    CONTRACT_SECTIONS,
)
from .styles import (
    apply_document_defaults,
    apply_body_style,
    apply_run_font,
    apply_title_style,
    apply_heading_style,
)


class ContractTemplateBuilder:
    """Строитель шаблона договора"""

    def __init__(self):
        self.doc = Document()
        apply_document_defaults(self.doc)

    def add_header(self):
        """Добавляет заголовок договора"""
        # Договор
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Договор")
        apply_title_style(run)
        p.paragraph_format.space_after = Pt(0)

        # возмездного оказания услуг
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("возмездного оказания услуг")
        apply_title_style(run)
        p.paragraph_format.space_after = Pt(0)

        # № {{contract_number}} от {{contract_date}}г.
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("№ {{contract_number}} от {{contract_date}}г.")
        apply_title_style(run)
        p.paragraph_format.space_after = Pt(12)

    def add_city_and_date(self):
        """Добавляет город и дату (г. Москва / дата)"""
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.allow_autofit = True

        # Устанавливаем ширину колонок
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(3)

        # Левая ячейка: г. Москва
        cell_left = table.rows[0].cells[0]
        p = cell_left.paragraphs[0]
        run = p.add_run("г. Москва")
        apply_run_font(run)

        # Правая ячейка: «{{day}}» {{date_text}} г.
        cell_right = table.rows[0].cells[1]
        p = cell_right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("«{{day}}» {{date_text}} г.")
        apply_run_font(run)

        # Пустая строка после
        self.doc.add_paragraph()

    def add_preamble(self):
        """Добавляет преамбулу с данными Исполнителя и Заказчика"""
        text = (
            f"{EXECUTOR_PREAMBLE}, и {{{{client_header}}}}, "
            "именуемый в дальнейшем «Заказчик», вместе именуемые «Стороны», "
            "а по отдельности «Сторона», заключили настоящий договор о нижеследующем:"
        )
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        apply_run_font(run)
        apply_body_style(p, first_line_indent=True)

    def add_section(self, number, title: str, paragraphs: list):
        """Добавляет раздел договора"""
        # Заголовок раздела
        heading = self.doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if number:
            heading_text = f"{number}. {title}"
        else:
            heading_text = title
        run = heading.add_run(heading_text)
        apply_heading_style(run)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.left_indent = Cm(1.25)

        # Пункты раздела
        for para_text in paragraphs:
            p = self.doc.add_paragraph()
            run = p.add_run(para_text)
            apply_run_font(run)
            apply_body_style(p, first_line_indent=True)

    def _add_requisites_table(
        self,
        with_heading: bool = False,
        add_empty_paragraph: bool = False
    ):
        """Добавляет таблицу реквизитов сторон"""
        if add_empty_paragraph:
            self.doc.add_paragraph()

        if with_heading:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run("7. Реквизиты Сторон:")
            apply_heading_style(run)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)

        # Таблица 5 строк x 2 колонки с рамками
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        table.autofit = True

        # Строка 1: Заголовки (Исполнитель | Заказчик)
        cell_left = table.rows[0].cells[0]
        cell_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell_left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run("Исполнитель")
        run.bold = True
        apply_run_font(run)

        cell_right = table.rows[0].cells[1]
        cell_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell_right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run("Заказчик")
        run.bold = True
        apply_run_font(run)

        # Строка 2: Основные данные
        cell_left = table.rows[1].cells[0]
        p = cell_left.paragraphs[0]
        run = p.add_run("{{executor_main}}")
        p.paragraph_format.space_before = Pt(6)
        apply_run_font(run)

        cell_right = table.rows[1].cells[1]
        p = cell_right.paragraphs[0]
        run = p.add_run("{{requisites_main}}")
        p.paragraph_format.space_before = Pt(6)
        apply_run_font(run)

        # Строка 3: "Банковские реквизиты:"
        cell_left = table.rows[2].cells[0]
        p = cell_left.paragraphs[0]
        run = p.add_run("Банковские реквизиты:")
        p.paragraph_format.space_after = Pt(6)
        apply_run_font(run)

        cell_right = table.rows[2].cells[1]
        p = cell_right.paragraphs[0]
        run = p.add_run("Банковские реквизиты:")
        p.paragraph_format.space_after = Pt(6)
        apply_run_font(run)

        # Строка 4: Банковские данные
        cell_left = table.rows[3].cells[0]
        p = cell_left.paragraphs[0]
        run = p.add_run("{{executor_bank}}")
        apply_run_font(run)

        cell_right = table.rows[3].cells[1]
        p = cell_right.paragraphs[0]
        run = p.add_run("{{requisites_bank}}")
        apply_run_font(run)

        # Строка 5: Подписи
        cell_left = table.rows[4].cells[0]
        p = cell_left.paragraphs[0]
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run(f"_________________/{EXECUTOR_DATA['name_short']} /")
        p.paragraph_format.space_after = Pt(6)
        apply_run_font(run)

        cell_right = table.rows[4].cells[1]
        p = cell_right.paragraphs[0]
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run("_________________/{{signatory}} /")
        p.paragraph_format.space_after = Pt(6)
        apply_run_font(run)

    def add_requisites_section(self):
        """Добавляет раздел 7: Реквизиты Сторон (таблица с рамками)"""
        self._add_requisites_table(with_heading=True)

    def add_page_break(self):
        """Добавляет разрыв страницы"""
        self.doc.add_page_break()

    def add_task_page(self):
        """Добавляет страницу 'Задание Заказчика № 1' с таблицей услуг"""
        # Заголовок
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Задание Заказчика № 1")
        apply_title_style(run)
        p.paragraph_format.space_after = Pt(0)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("возмездного оказания услуг")
        apply_title_style(run)
        p.paragraph_format.space_after = Pt(0)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("{{contract_number}} от {{contract_date}}г.")
        apply_title_style(run)
        p.paragraph_format.space_after = Pt(12)

        # Текст введения
        p = self.doc.add_paragraph()
        run = p.add_run(
            "На основании Договора возмездного оказания услуг от «{{day}}» {{date_text}} г. "
            "Исполнитель обязуется:\n"
        )
        apply_run_font(run)
        apply_body_style(p, first_line_indent=False)

        # Таблица услуг (3 колонки: №, Наименование, Стоимость)
        table = self.doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.autofit = False

        # Ширина колонок: 1см + равные части для остальных
        col_widths = [Cm(1), Cm(8), Cm(8)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

        # Заголовки (первая колонка без заголовка - для номера)
        headers = ["", "Наименование услуги", "Стоимость (руб.) и порядок оплаты"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            if header:
                run = p.add_run(header)
                apply_run_font(run)

        # Маркер для заполнения - используем вторую строку
        cell = table.rows[1].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run("{{services}}")
        apply_run_font(run)

        # Пустые ячейки справа
        for col_idx in [1, 2]:
            cell = table.rows[1].cells[col_idx]
            p = cell.paragraphs[0]
            p.add_run("")

    def add_task_requisites(self):
        """Добавляет блок реквизитов и подписей на странице задания"""
        self._add_requisites_table(with_heading=False, add_empty_paragraph=True)

    def build(self) -> Document:
        """Собирает полный шаблон договора"""
        self.add_header()
        self.add_city_and_date()
        self.add_preamble()

        # Добавляем все разделы из CONTRACT_SECTIONS
        for section in CONTRACT_SECTIONS:
            self.add_section(
                section['number'],
                section['title'],
                section['paragraphs']
            )

        self.add_requisites_section()
        self.add_page_break()
        self.add_task_page()
        self.add_task_requisites()

        return self.doc

    def save(self, path: str):
        """Сохраняет шаблон в файл"""
        self.build().save(path)

    def to_bytes(self) -> bytes:
        """Возвращает шаблон как bytes"""
        buffer = BytesIO()
        self.build().save(buffer)
        return buffer.getvalue()


def generate_template() -> Document:
    """Фабричная функция для создания шаблона"""
    builder = ContractTemplateBuilder()
    return builder.build()


def generate_template_file(output_path: str = None):
    """Генерирует и сохраняет шаблон в файл"""
    if output_path is None:
        output_path = str(TEMPLATE_PATH)
    builder = ContractTemplateBuilder()
    builder.save(output_path)
    return output_path
