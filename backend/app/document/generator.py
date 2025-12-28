from io import BytesIO
from decimal import Decimal
from docx import Document
from docx.shared import Pt

from app.models import Contract, CLIENT_TYPES
from .constants import TEMPLATE_PATH
from .tables import fill_services_table
from .replacements import replace_in_paragraph, build_replacements, get_full_name, build_requisites
from .utils import number_to_words_ru


def generate_contract_document(contract: Contract) -> bytes:
    if not TEMPLATE_PATH.exists():
        return generate_fallback_document(contract)

    doc = Document(TEMPLATE_PATH)
    total = fill_services_table(doc, contract.services)
    replacements = build_replacements(contract, total)

    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_fallback_document(contract: Contract) -> bytes:
    doc = Document()
    doc.add_heading(f"Договор № {contract.number}", 0)
    doc.add_paragraph(f"Дата: {contract.date.strftime('%d.%m.%Y')}")
    doc.add_paragraph("")

    doc.add_heading("Клиент", level=1)
    client = contract.client
    ct = client.client_type or "ip"
    type_label = CLIENT_TYPES.get(ct, "")

    doc.add_paragraph(f"Тип: {type_label}")
    doc.add_paragraph(f"Наименование: {client.name}")
    doc.add_paragraph(f"ФИО: {get_full_name(client)}")
    if client.inn:
        doc.add_paragraph(f"ИНН: {client.inn}")
    if client.ogrn:
        ogrn_label = "ОГРНИП" if ct == "ip" else "ОГРН"
        doc.add_paragraph(f"{ogrn_label}: {client.ogrn}")
    if client.kpp:
        doc.add_paragraph(f"КПП: {client.kpp}")
    if client.address:
        doc.add_paragraph(f"Адрес: {client.address}")
    if client.phone:
        doc.add_paragraph(f"Телефон: {client.phone}")
    if client.email:
        doc.add_paragraph(f"Email: {client.email}")
    if client.settlement_account:
        doc.add_paragraph(f"Расчётный счёт: {client.settlement_account}")

    if client.bank:
        doc.add_paragraph(f"Банк: {client.bank.name}")
        doc.add_paragraph(f"БИК: {client.bank.bik}")
        doc.add_paragraph(f"Кор. счёт: {client.bank.correspondent_account}")

    doc.add_paragraph("")
    doc.add_heading("Услуги", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Наименование"
    hdr_cells[2].text = "Стоимость"
    hdr_cells[3].text = "Порядок оплаты"

    total = Decimal("0")
    for i, service in enumerate(contract.services, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = service.name
        row_cells[2].text = f"{service.price:,.2f} руб."
        row_cells[3].text = service.payment_terms
        total += service.price

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
