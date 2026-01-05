from decimal import Decimal

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .utils import number_to_words_ru
from .styles import apply_run_font, set_cell_width

# Ширина колонок таблицы услуг
SERVICE_TABLE_COL_WIDTHS = [Cm(1), Cm(8), Cm(8)]


def set_cell_multiline_text(cell, text: str):
    """Устанавливает многострочный текст с реальными параграфами.

    В отличие от cell.text = "line1\\nline2", создает настоящие
    параграфы, которые корректно отображаются в LibreOffice.
    """
    lines = text.split('\n')
    # Очищаем первый параграф
    cell.paragraphs[0].clear()

    for i, line in enumerate(lines):
        if i == 0:
            run = cell.paragraphs[0].add_run(line)
        else:
            p = cell.add_paragraph()
            run = p.add_run(line)
        apply_run_font(run)


def find_services_table(doc):
    """Находит таблицу с маркером {{services}}"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{services}}" in cell.text:
                    return table, row
    return None, None


def fill_services_table(doc, services: list) -> Decimal:
    """Заполняет таблицу услуг, возвращает итоговую сумму"""
    table, marker_row = find_services_table(doc)
    if not table:
        return Decimal("0")

    table._tbl.remove(marker_row._tr)

    total = Decimal("0")
    for i, service in enumerate(services, 1):
        row = table.add_row()
        # Применяем ширину колонок к новой строке (LibreOffice)
        for idx, width in enumerate(SERVICE_TABLE_COL_WIDTHS):
            set_cell_width(row.cells[idx], width.twips)

        # Номер услуги
        row.cells[0].paragraphs[0].clear()
        run = row.cells[0].paragraphs[0].add_run(f"{i}.")
        apply_run_font(run)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Название услуги
        row.cells[1].paragraphs[0].clear()
        run = row.cells[1].paragraphs[0].add_run(service.name)
        apply_run_font(run)
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Стоимость и порядок оплаты (многострочный текст)
        price_text = f"Стоимость: {service.price:,.0f} руб. ({number_to_words_ru(service.price)}).\nПорядок оплаты:\n{service.payment_terms}"
        set_cell_multiline_text(row.cells[2], price_text)
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[2].paragraphs[0].paragraph_format.space_before = Pt(6)
        row.cells[2].paragraphs[0].paragraph_format.space_after = Pt(6)

        total += service.price

    return total
